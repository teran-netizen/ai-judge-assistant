"""
Сервис генерации .docx документов из текста решения.

Форматирование: Times New Roman 12pt, поля 2 cm, заголовок по центру.
"""
import io
import re
import logging

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def _strip_markdown(text: str) -> str:
    """Убирает markdown-разметку из текста (страховка от DeepSeek)."""
    # **жирный** и __жирный__
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    # *курсив* и _курсив_ (но не подчёркивания в словах типа case_id)
    text = re.sub(r'(?<!\w)\*(.+?)\*(?!\w)', r'\1', text)
    # # заголовки (строки начинающиеся с #)
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    # --- горизонтальные линии
    text = re.sub(r'^-{3,}$', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\*{3,}$', '', text, flags=re.MULTILINE)
    text = re.sub(r'^_{3,}$', '', text, flags=re.MULTILINE)
    return text


def build_docx(title: str, text: str) -> io.BytesIO:
    """
    Генерирует .docx файл из заголовка и текста решения.

    Args:
        title: Заголовок документа (название дела).
        text: Полный текст решения (абзацы разделены \\n).

    Returns:
        BytesIO буфер с готовым .docx файлом (seek(0) уже выполнен).
    """
    doc = Document()

    # Поля страницы — 2 cm со всех сторон (стандарт судебных документов)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Заголовок по центру
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Убираем markdown-разметку (страховка) и разбиваем на абзацы
    text = _strip_markdown(text)
    paragraphs = text.split("\n")
    for para_text in paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(para_text)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        # Кириллица: устанавливаем шрифт для Complex Script и East Asian
        r_elem = run._element
        rPr = r_elem.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:cs'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Очистка метаданных — никаких следов генерации
    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.comments = ""
    cp.keywords = ""
    cp.category = ""
    cp.subject = ""
    cp.title = ""
    cp.revision = 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    logger.info("DOCX generated: title=%r, paragraphs=%d, size=%d bytes",
                title, len(paragraphs), buffer.getbuffer().nbytes)
    return buffer


def safe_filename(title: str, case_id: str) -> str:
    """
    Генерирует безопасное имя файла для скачивания.
    
    Формат: {название}_{дата}_{id}.docx
    Пример: Исковое_заявление_о_взыскании_2026-05-05_a1b2c3d4.docx
    """
    from datetime import datetime
    clean = title or ""
    safe_title = re.sub(r'[^\w\s\-Ѐ-ӿ]', '', clean)[:40].strip()
    if not safe_title or len(safe_title) < 3:
        safe_title = f"документ_{case_id[:8]}"
    date_str = datetime.now().strftime("%Y-%m-%d")
    short_id = case_id[:8]
    return f"{safe_title}_{date_str}_{short_id}.docx"
